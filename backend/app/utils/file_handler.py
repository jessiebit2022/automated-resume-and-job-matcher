import os
import logging
from typing import Optional
import PyPDF2
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)

class FileHandler:
    """Handle file operations and text extraction from various formats."""
    
    def extract_text(self, file_path: str, file_extension: str) -> Optional[str]:
        """Extract text from uploaded file based on its extension."""
        try:
            if file_extension == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['doc', 'docx']:
                return self._extract_from_docx(file_path)
            else:
                logger.error(f"Unsupported file extension: {file_extension}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    def _extract_from_pdf(self, file_path: str) -> Optional[str]:
        """Extract text from PDF file using pdfplumber (more reliable than PyPDF2)."""
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\\n"
            
            if not text.strip():
                # Fallback to PyPDF2 if pdfplumber fails
                logger.info("pdfplumber extraction failed, trying PyPDF2...")
                return self._extract_from_pdf_pypdf2(file_path)
                
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting PDF with pdfplumber: {str(e)}")
            # Fallback to PyPDF2
            return self._extract_from_pdf_pypdf2(file_path)
    
    def _extract_from_pdf_pypdf2(self, file_path: str) -> Optional[str]:
        """Fallback PDF extraction using PyPDF2."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\\n"
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting PDF with PyPDF2: {str(e)}")
            return None
    
    def _extract_from_docx(self, file_path: str) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting DOCX: {str(e)}")
            return None
    
    def validate_file(self, file_path: str) -> bool:
        """Validate that file exists and is readable."""
        try:
            return os.path.exists(file_path) and os.path.isfile(file_path) and os.access(file_path, os.R_OK)
        except Exception:
            return False
    
    def get_file_size(self, file_path: str) -> int:
        """Get file size in bytes."""
        try:
            return os.path.getsize(file_path)
        except Exception:
            return 0